import json
import re
import time
import zipfile
from io import BytesIO
from pathlib import Path
from urllib.parse import quote

from watchfiles import awatch

from docx import Document
from fastapi import Depends, FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, StreamingResponse
from pydantic import BaseModel
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

from app.domains.writing.expert_agent.services.character import CharacterService
from app.domains.writing.meta import MetaAgentService
from app.platform.core.settings import get_settings
from app.create_type.store import CreateTypeStore
from app.create_type.optimizer import StyleOptimizer
from app.create_type.router import init_style_module, router as style_router
from app.platform.state.thread_store import ThreadStore
from app.platform.trace import TraceRecorder
from app.schemas.character import CharacterGenerateRequest, CharacterGenerateResponse
from app.schemas.screenplay import (
    InitResponse,
    ScreenplayGenerateRequest,
    ScreenplayGenerateResponse,
    StorylineGraphEvent,
    StorylineGraphStoryline,
    ThreadCreateRequest,
    ThreadSummary,
    ThreadUpdateRequest,
    WorkspaceBootstrapResponse,
    WorkspaceCharacterContent,
    WorkspaceCreateRequest,
    WorkspaceDetailOutlineContent,
    WorkspaceNovelChaptersContent,
    WorkspaceOutlineContent,
    WorkspaceStorylineContent,
    WorkspaceStorylineGraphContent,
    WorkspaceWorldviewContent,
    WorkspaceSummary,
)
from app.schemas.checkpoint import CheckpointState
from app.platform.trace import TraceDetail, TraceRunSummary
from app.platform.agent.middleware import TraceCallbackHandler
from app.auth import auth_router, current_user, CurrentUser
from app.auth.bootstrap import bootstrap_admin
from app.admin import me_router, admin_router
from app.platform.core.checkpoint_pool import CheckpointPool, init_checkpoint_pool, get_checkpoint_pool
from app.platform.core.security import load_master_key
from app.platform.core.db import Database, init_database, get_database, UserRepository


_active_generations = 0


def _log(event: str, **fields) -> None:
    """诊断日志：结构化 JSON 到 stdout（flush 保证即时输出）。

    为 Phase 1 根因诊断服务——定位 SSE 连接泄漏 / 请求挂起 / 锁竞争。
    设计文档 §5 约定：轻量、零依赖、JSON 到 stdout。
    """
    record = {"ts": time.strftime("%Y-%m-%dT%H:%M:%S"), "event": event, **fields}
    print(json.dumps(record, ensure_ascii=False), flush=True)


def _markdown_to_plain_text(markdown: str) -> str:
    lines = []
    for line in markdown.splitlines():
        stripped = line.strip()
        if stripped.startswith("## "):
            lines.append(stripped[3:].strip())
        elif stripped.startswith("# "):
            lines.append(stripped[2:].strip())
        elif stripped == "---":
            lines.append("")
        else:
            lines.append(line)
    return "\n".join(lines).strip()


def _escape_pdf_text(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _safe_download_name(name: str, fallback: str, max_length: int = 80) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "", name).strip().strip(".")
    return (cleaned or fallback)[:max_length]


def _build_novel_docx(markdown: str, title: str) -> bytes:
    document = Document()
    document.add_heading(title, level=1)
    for block in _markdown_to_plain_text(markdown).split("\n\n"):
        text = block.strip()
        if not text or text == title:
            continue
        document.add_paragraph(text)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_novel_docx_zip(content: WorkspaceNovelChaptersContent) -> bytes:
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for index, chapter in enumerate(content.chapters, start=1):
            title = chapter.title.strip() or Path(chapter.filename).stem or f"chapter-{index:03d}"
            safe_title = _safe_download_name(title, f"chapter-{index:03d}")
            archive.writestr(f"{index:03d}-{safe_title}.docx", _build_novel_docx(chapter.markdown, title))
    return buffer.getvalue()


def _build_novel_pdf(markdown: str, title: str) -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54,
        title=title,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "NovelTitle",
        parent=styles["Title"],
        fontName="STSong-Light",
        fontSize=20,
        leading=28,
        spaceAfter=24,
    )
    chapter_style = ParagraphStyle(
        "ChapterTitle",
        parent=styles["Heading2"],
        fontName="STSong-Light",
        fontSize=15,
        leading=22,
        spaceBefore=8,
        spaceAfter=12,
    )
    body_style = ParagraphStyle(
        "NovelBody",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=11,
        leading=19,
        firstLineIndent=22,
        spaceAfter=7,
    )

    story = [Paragraph(_escape_pdf_text(title), title_style)]
    first_chapter = True
    for block in _markdown_to_plain_text(markdown).split("\n\n"):
        text = block.strip()
        if not text:
            continue
        if text.startswith("第") and "章" in text[:8]:
            if not first_chapter:
                story.append(PageBreak())
            first_chapter = False
            story.append(Paragraph(_escape_pdf_text(text), chapter_style))
            continue
        story.append(Paragraph(_escape_pdf_text(text).replace("\n", "<br/>"), body_style))
        story.append(Spacer(1, 4))

    doc.build(story)
    return buffer.getvalue()


async def _event_generator(payload: ScreenplayGenerateRequest, thread: ThreadSummary, *, owner_id: str | None = None):
    """SSE 生成器已迁到 routers/screenplay.py（PR-14）。保留此 shim 供过渡期调用。"""
    from app.routers.screenplay import _event_generator as _gen
    async for chunk in _gen(payload, thread, owner_id=owner_id):
        yield chunk


settings = get_settings()
workspace_root = Path(__file__).resolve().parents[1] / settings.workspace_root
trace_recorder = TraceRecorder()

# 多用户地基：元数据库 + checkpoint 分库池
_master_key = load_master_key(settings.master_key)
_database = Database(settings.db_path, _master_key)
init_database(_database)

_checkpoints_root = Path(__file__).resolve().parents[1] / "checkpoints"
_checkpoint_pool = CheckpointPool(_checkpoints_root)
init_checkpoint_pool(_checkpoint_pool)

thread_store = ThreadStore(_database, workspace_root)
style_store = CreateTypeStore(_database, thread_store.workspaces)
style_optimizer = StyleOptimizer(settings)
init_style_module(style_store, style_optimizer)

# 全局 checkpointer：仅作管理员兜底与同步 delete_thread 路径用；
# 普通用户走 CheckpointPool 分库（services 内部按 owner_id 解析）。
from langgraph.checkpoint.sqlite.aio import AsyncSqliteSaver

_global_checkpoint_db_path = workspace_root.parent / "checkpoints.db"
_checkpointer_cm = AsyncSqliteSaver.from_conn_string(str(_global_checkpoint_db_path))

agent_service: MetaAgentService | None = None
character_service: CharacterService | None = None
image_agent_service = None  # ImageAgentService 实例（Phase 3）


async def _lifespan(application: FastAPI):
    global agent_service, character_service, image_agent_service
    checkpointer = await _checkpointer_cm.__aenter__()
    if agent_service is None:
        agent_service = MetaAgentService(settings, workspace_root, trace_recorder, style_store, checkpointer)
    if character_service is None:
        character_service = CharacterService(settings, workspace_root, trace_recorder, checkpointer)
    if image_agent_service is None:
        from app.domains.image.agent import ImageAgentService
        image_agent_service = ImageAgentService(settings, workspace_root, trace_recorder, checkpointer)
        from app.domains.image.router import init_image_routes
        init_image_routes(image_agent_service, thread_store, trace_recorder)
    # PR-14：注入 router 共享 context（screenplay/character router 通过 get_*() 访问）
    from app.routers.context import init_router_context
    init_router_context(
        thread_store=thread_store, style_store=style_store, trace_recorder=trace_recorder,
        agent_service=agent_service, character_service=character_service,
        image_agent_service=image_agent_service, style_optimizer=style_optimizer,
    )
    # 多用户：引导管理员账号（幂等）
    bootstrap_admin()
    # 启动 trace 写盘 drain 协程：append_event 不再同步写盘，改入内存缓冲，
    # 由此后台协程成批落盘（to_thread，不占事件循环）。
    trace_recorder.start_drain()
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    yield
    # 关闭 drain 并刷掉残余事件，保证进程退出前 trace 数据完整。
    await trace_recorder.aclose()
    await _checkpointer_cm.__aexit__(None, None, None)
    await _checkpoint_pool.aclose_all()


app = FastAPI(
    title="Writer Agent API",
    version="0.1.0",
    description="Minimal screenplay generation backend built with DeepAgents and FastAPI.",
    lifespan=_lifespan,
)
app.add_middleware(
    CORSMiddleware,
    allow_origins=[settings.writer_frontend_origin],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
app.include_router(style_router)
app.include_router(auth_router)
app.include_router(me_router)
app.include_router(admin_router)
# image domain 路由（Phase 3：图片服务端点 + Skill 管理）
from app.domains.image.router import router as image_router
app.include_router(image_router)
# PR-14：生成端点 router（从 main.py 抽出，路径前缀 /api）
from app.routers.screenplay import router as screenplay_router
from app.routers.character import router as character_router
app.include_router(screenplay_router, prefix="/api")
app.include_router(character_router, prefix="/api")


@app.middleware("http")
async def _log_http(request: Request, call_next):
    """记录每个 HTTP 请求的耗时与状态码。

    注意：对 SSE（StreamingResponse），Starlette 中间件的 ms ≈ 首字节时间，
    非连接全程时长——后者由 _event_generator / _workspace_watch_generator
    内部的 sse_close/sse_error 埋点覆盖。中间件主职是普通短请求（fetchInit 等）的耗时。
    """
    start = time.perf_counter()
    status_code = 0
    try:
        response = await call_next(request)
        status_code = response.status_code
        return response
    finally:
        _log("http", method=request.method, path=request.url.path,
             status=status_code, ms=int((time.perf_counter() - start) * 1000))


@app.get("/health")
def healthcheck() -> dict[str, str]:
    return {"status": "ok", "mode": settings.writer_agent_mode}


@app.get("/api/init", response_model=InitResponse)
def init_page(user: CurrentUser = Depends(current_user)) -> InitResponse:
    """页面首次加载：一次返回当前用户的 workspaces + styles。"""
    return InitResponse(
        workspaces=thread_store.list_workspaces(user.user_id),
        styles=style_store.list_styles(user.user_id),
    )


@app.get("/api/workspaces/{workspace_id}/bootstrap", response_model=WorkspaceBootstrapResponse)
def bootstrap_workspace(workspace_id: str, user: CurrentUser = Depends(current_user)) -> WorkspaceBootstrapResponse:
    """选中工作区后：一次返回 threads + 全部面板内容，替代 5 个独立请求。"""
    data = thread_store.bootstrap_workspace(user.user_id, workspace_id)
    if data is None:
        raise HTTPException(status_code=404, detail="Workspace not found")

    return WorkspaceBootstrapResponse(
        threads=data["threads"],
        outline=data["outline"],
        storyline=data["storyline"],
        detail_outline=data["detail_outline"],
        characters=data["characters"],
        novel=data["novel"],
        worldview=data["worldview"],
    )


@app.get("/api/workspaces", response_model=list[WorkspaceSummary])
def list_workspaces(user: CurrentUser = Depends(current_user)) -> list[WorkspaceSummary]:
    return thread_store.list_workspaces(user.user_id)


@app.post("/api/workspaces", response_model=WorkspaceSummary)
def create_workspace(payload: WorkspaceCreateRequest, user: CurrentUser = Depends(current_user)) -> WorkspaceSummary:
    # 配额检查（T2.8）
    users = UserRepository(get_database())
    owner = users.get_by_id(user.user_id)
    quota = owner["workspace_quota"] if owner else settings.default_workspace_quota
    if users.workspace_count(user.user_id) >= quota:
        raise HTTPException(
            status_code=409,
            detail=f"作品数量已达上限（{quota} 部）",
        )
    try:
        return thread_store.create_workspace(user.user_id, payload.title, payload.domain)
    except FileExistsError as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc


@app.get("/api/workspaces/{workspace_id}/outline", response_model=WorkspaceOutlineContent)
def get_workspace_outline(workspace_id: str, user: CurrentUser = Depends(current_user)) -> WorkspaceOutlineContent:
    content = thread_store.artifacts.read_workspace_outline(user.user_id, workspace_id)
    if content is None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    return content


@app.get("/api/workspaces/{workspace_id}/storyline", response_model=WorkspaceStorylineContent)
def get_workspace_storyline(workspace_id: str, user: CurrentUser = Depends(current_user)) -> WorkspaceStorylineContent:
    content = thread_store.artifacts.read_workspace_storyline(user.user_id, workspace_id)
    if content is None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    return content


@app.get("/api/workspaces/{workspace_id}/detail-outline", response_model=WorkspaceDetailOutlineContent)
def get_workspace_detail_outline(workspace_id: str, user: CurrentUser = Depends(current_user)) -> WorkspaceDetailOutlineContent:
    content = thread_store.artifacts.read_workspace_detail_outline(user.user_id, workspace_id)
    if content is None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    return content


@app.get("/api/workspaces/{workspace_id}/worldview", response_model=WorkspaceWorldviewContent)
def get_workspace_worldview(workspace_id: str, user: CurrentUser = Depends(current_user)) -> WorkspaceWorldviewContent:
    content = thread_store.artifacts.read_workspace_worldview(user.user_id, workspace_id)
    if content is None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    return content


@app.get("/api/workspaces/{workspace_id}/storyline-graph", response_model=WorkspaceStorylineGraphContent)
def get_workspace_storyline_graph(workspace_id: str, user: CurrentUser = Depends(current_user)) -> WorkspaceStorylineGraphContent:
    """故事线流程图（竖向泳道时间轴）。读取时按需生成兜底——图缺失/过期自动重生成。

    生成逻辑在 API 层（而非 thread_store）：避免 core（基础设施层）反向依赖 writer。
    thread_store 只负责读 markdown，结构化数据（events/storylines/t_map）在此解析填充。
    """
    content = thread_store.artifacts.read_workspace_storyline_graph(user.user_id, workspace_id)
    if content is None:
        raise HTTPException(status_code=404, detail="Workspace not found")

    # 按需生成兜底：图缺失/过期 → 重生成（storyline_graph 是派生视图，幂等安全）
    from app.domains.writing.expert_agent.services.storyline_graph import (
        build_storyline_graph_data,
        generate_storyline_graph,
        is_stale,
    )
    workspace = thread_store.get_workspace(user.user_id, workspace_id)
    if workspace is not None:
        ws_path = Path(workspace.workspace_path)
        stale = is_stale(ws_path)
        if stale:
            generate_storyline_graph(ws_path)
            content = thread_store.artifacts.read_workspace_storyline_graph(user.user_id, workspace_id)
            if content is None:
                raise HTTPException(status_code=404, detail="Workspace not found")
            content.stale = True
        # 填充结构化数据（reactflow 自定义布局用）
        data = build_storyline_graph_data(ws_path)
        if data is not None:
            content.storylines = [
                StorylineGraphStoryline(
                    id=sl.id, name=sl.name, type=sl.type, status=sl.status,
                    direction=sl.direction, key_events=sl.key_events,
                ) for sl in data.storylines
            ]
            content.events = {
                eid: StorylineGraphEvent(
                    id=ev.id, name=ev.name, type=ev.type,
                    storylines=list(ev.storylines), group=ev.group, doc_order=ev.doc_order,
                ) for eid, ev in data.events.items()
            }
            content.t_map = dict(data.t_map)
            content.storyline_count = len(data.storylines)
            content.event_count = len(data.events)
    return content


@app.get("/api/workspaces/{workspace_id}/characters", response_model=WorkspaceCharacterContent)
def get_workspace_characters(workspace_id: str, user: CurrentUser = Depends(current_user)) -> WorkspaceCharacterContent:
    content = thread_store.artifacts.read_workspace_characters(user.user_id, workspace_id)
    if content is None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    return content


@app.get("/api/workspaces/{workspace_id}/novel", response_model=WorkspaceNovelChaptersContent)
def get_workspace_novel(workspace_id: str, user: CurrentUser = Depends(current_user)) -> WorkspaceNovelChaptersContent:
    content = thread_store.artifacts.read_workspace_novel_chapters(user.user_id, workspace_id)
    if content is None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    return content


def _sse_event(event_type: str, payload: object) -> str:
    data = json.dumps(payload, ensure_ascii=False, default=str)
    return f"event: {event_type}\ndata: {data}\n\n"


def _classify_changes(changes, workspace_path: Path) -> set[str]:
    categories: set[str] = set()
    for _change_type, path_str in changes:
        try:
            rel = Path(path_str).relative_to(workspace_path)
        except ValueError:
            continue
        parts = rel.parts
        if not parts:
            continue
        top = parts[0]
        if top in ("outline.md", "evaluation.md"):
            categories.add("outline")
        elif top == "storyline.md" or (len(parts) > 1 and parts[0] == "storyline"):
            categories.add("storyline")
        elif top == "worldview.md":
            categories.add("worldview")
        elif len(parts) > 1 and parts[0] == "detail":
            categories.add("detail_outline")
        elif len(parts) > 1 and parts[0] == "character":
            categories.add("characters")
        elif len(parts) > 1 and parts[0] == "chapter":
            categories.add("novel")
        elif top == "novel.md":
            categories.add("novel")
    return categories


async def _workspace_watch_generator(owner_id: str, workspace_id: str, workspace_path: Path):
    _log("sse_open", channel="watch", workspace_id=workspace_id)
    start = time.perf_counter()
    try:
        async for changes in awatch(
            workspace_path,
            watch_filter=lambda _change, path: Path(path).suffix == ".md",
            debounce=400,
            step=50,
            recursive=True,
            ignore_permission_denied=True,
        ):
            categories = _classify_changes(changes, workspace_path)
            if not categories:
                continue
            if "outline" in categories:
                content = thread_store.artifacts.read_workspace_outline(owner_id, workspace_id)
                if content is not None:
                    yield _sse_event("outline", content.model_dump())
            if "storyline" in categories:
                content = thread_store.artifacts.read_workspace_storyline(owner_id, workspace_id)
                if content is not None:
                    yield _sse_event("storyline", content.model_dump())
            if "worldview" in categories:
                content = thread_store.artifacts.read_workspace_worldview(owner_id, workspace_id)
                if content is not None:
                    yield _sse_event("worldview", content.model_dump())
            if "detail_outline" in categories:
                content = thread_store.artifacts.read_workspace_detail_outline(owner_id, workspace_id)
                if content is not None:
                    yield _sse_event("detail_outline", content.model_dump())
            if "characters" in categories:
                content = thread_store.artifacts.read_workspace_characters(owner_id, workspace_id)
                if content is not None:
                    yield _sse_event("characters", content.model_dump())
            if "novel" in categories:
                content = thread_store.artifacts.read_workspace_novel_chapters(owner_id, workspace_id)
                if content is not None:
                    yield _sse_event("novel", content.model_dump())
        _log("sse_close", channel="watch", workspace_id=workspace_id,
             ms=int((time.perf_counter() - start) * 1000))
    except BaseException as exc:
        _log("sse_error", channel="watch", workspace_id=workspace_id,
             error=type(exc).__name__, ms=int((time.perf_counter() - start) * 1000))
        raise


@app.get("/api/workspaces/{workspace_id}/watch")
async def watch_workspace(workspace_id: str, user: CurrentUser = Depends(current_user)):
    workspace = thread_store.get_workspace(user.user_id, workspace_id)
    if workspace is None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    workspace_path = Path(workspace.workspace_path)
    if not workspace_path.exists():
        raise HTTPException(status_code=404, detail="Workspace directory missing")
    return StreamingResponse(
        _workspace_watch_generator(user.user_id, workspace_id, workspace_path),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.get("/api/workspaces/{workspace_id}/novel/export.pdf")
def export_workspace_novel_pdf(workspace_id: str, user: CurrentUser = Depends(current_user)) -> Response:
    workspace = thread_store.get_workspace(user.user_id, workspace_id)
    if workspace is None:
        raise HTTPException(status_code=404, detail="Workspace not found")

    content = thread_store.artifacts.read_workspace_novel(user.user_id, workspace_id)
    if content is None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    if not content.markdown.strip():
        raise HTTPException(status_code=404, detail="Novel content not found")

    filename = f"{workspace.title or workspace_id}.pdf"
    pdf = _build_novel_pdf(content.markdown, workspace.title or "小说正文")
    return Response(
        content=pdf,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"},
    )


@app.get("/api/workspaces/{workspace_id}/novel/export-word.zip")
def export_workspace_novel_word_zip(workspace_id: str, user: CurrentUser = Depends(current_user)) -> Response:
    workspace = thread_store.get_workspace(user.user_id, workspace_id)
    if workspace is None:
        raise HTTPException(status_code=404, detail="Workspace not found")

    content = thread_store.artifacts.read_workspace_novel_chapters(user.user_id, workspace_id)
    if content is None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    if not content.chapters:
        raise HTTPException(status_code=404, detail="Novel content not found")

    filename_base = _safe_download_name(workspace.title or workspace_id, workspace_id)
    archive = _build_novel_docx_zip(content)
    return Response(
        content=archive,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(f'{filename_base}-word.zip')}"},
    )


@app.delete("/api/workspaces/{workspace_id}")
async def delete_workspace(workspace_id: str, user: CurrentUser = Depends(current_user)) -> dict[str, str | bool | list[str]]:
    try:
        deleted_thread_ids = thread_store.delete_workspace(user.user_id, workspace_id)
    except FileNotFoundError as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    if deleted_thread_ids is None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    # T2.9：清理 checkpoint（分库）+ trace
    for thread_id in deleted_thread_ids:
        await agent_service.delete_thread_checkpoint(thread_id, owner_id=user.user_id)
        await character_service.delete_thread_checkpoint(thread_id)
    return {"status": "ok", "deleted": workspace_id, "deleted_threads": deleted_thread_ids}


@app.get("/api/threads", response_model=list[ThreadSummary])
def list_threads(workspace_id: str | None = None, user: CurrentUser = Depends(current_user)) -> list[ThreadSummary]:
    return thread_store.list_threads(user.user_id, workspace_id)


@app.post("/api/threads", response_model=ThreadSummary)
def create_thread(payload: ThreadCreateRequest, user: CurrentUser = Depends(current_user)) -> ThreadSummary:
    try:
        return thread_store.create_thread(user.user_id, payload.workspace_id, payload.session_name)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc


@app.patch("/api/threads/{thread_id}", response_model=ThreadSummary)
def update_thread(thread_id: str, payload: ThreadUpdateRequest, user: CurrentUser = Depends(current_user)) -> ThreadSummary:
    try:
        thread = thread_store.update_thread_name(user.user_id, thread_id, payload.session_name)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    return thread


@app.delete("/api/threads/{thread_id}")
async def delete_thread(thread_id: str, user: CurrentUser = Depends(current_user)) -> dict[str, str | bool]:
    thread = thread_store.get_thread(user.user_id, thread_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    try:
        trace_recorder.delete_thread_runs(thread)
    except ValueError as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc

    deleted = thread_store.delete_thread(user.user_id, thread_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="Thread not found")
    await agent_service.delete_thread_checkpoint(thread_id, owner_id=user.user_id)
    await character_service.delete_thread_checkpoint(thread_id)
    return {"status": "ok", "deleted": thread_id}


@app.get("/api/threads/{thread_id}/outline", response_model=WorkspaceOutlineContent)
def get_thread_outline(thread_id: str, user: CurrentUser = Depends(current_user)) -> WorkspaceOutlineContent:
    content = thread_store.artifacts.read_thread_outline(user.user_id, thread_id)
    if content is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    return content


@app.get("/api/threads/{thread_id}/checkpoint", response_model=CheckpointState)
async def get_thread_checkpoint(thread_id: str, user: CurrentUser = Depends(current_user)) -> CheckpointState:
    thread = thread_store.get_thread(user.user_id, thread_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    return await agent_service.get_thread_checkpoint(thread_id, owner_id=user.user_id)


@app.get("/api/threads/{thread_id}/traces", response_model=list[TraceRunSummary])
def list_thread_traces(thread_id: str, user: CurrentUser = Depends(current_user)) -> list[TraceRunSummary]:
    thread = thread_store.get_thread(user.user_id, thread_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    return trace_recorder.list_runs(thread)


@app.get("/api/threads/{thread_id}/traces/{trace_id}", response_model=TraceDetail)
def get_thread_trace(thread_id: str, trace_id: str, user: CurrentUser = Depends(current_user)) -> TraceDetail:
    thread = thread_store.get_thread(user.user_id, thread_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    try:
        detail = trace_recorder.read_run(thread, trace_id)
    except FileNotFoundError as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    if detail is None:
        raise HTTPException(status_code=404, detail="Trace not found")
    return detail


@app.delete("/api/threads/{thread_id}/traces/{trace_id}")
def delete_thread_trace(thread_id: str, trace_id: str, user: CurrentUser = Depends(current_user)) -> dict[str, str]:
    thread = thread_store.get_thread(user.user_id, thread_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    try:
        deleted = trace_recorder.delete_run(thread, trace_id)
    except ValueError as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc
    if not deleted:
        raise HTTPException(status_code=404, detail="Trace not found")
    return {"status": "ok", "deleted": trace_id}


class ImageGenerateRequest(BaseModel):
    """文生图生成请求（已迁到 domains/image/router.py，此 shim 供 main 内部引用）。"""
    thread_id: str
    prompt: str
    trace_id: str | None = None
    resume: dict | None = None
    selected_skill_ids: list[str] | None = None


