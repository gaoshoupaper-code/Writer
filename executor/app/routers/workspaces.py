"""workspaces 路由（PR-14 从 main.py 抽出）。

工作区 CRUD + 产物读取（outline/storyline/detail/worldview/graph/characters/novel）
+ SSE watch + PDF/DOCX 导出。

辅助函数（导出格式化 + watch 分类）随端点迁入此模块。
"""

from __future__ import annotations

import json
import re
import time
import zipfile
from io import BytesIO
from pathlib import Path
from urllib.parse import quote

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response, StreamingResponse
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer
from watchfiles import awatch
from docx import Document

from app.auth import CurrentUser, current_user
from app.routers.context import _log, get_agent_service, get_character_service, get_thread_store
from app.schemas.screenplay import (
    StorylineGraphEvent,
    StorylineGraphStoryline,
    WorkspaceBootstrapResponse,
    WorkspaceCharacterContent,
    WorkspaceCreateRequest,
    WorkspaceDetailOutlineContent,
    WorkspaceNovelChaptersContent,
    WorkspaceNovelContent,
    WorkspaceOutlineContent,
    WorkspaceStorylineContent,
    WorkspaceStorylineGraphContent,
    WorkspaceSummary,
    WorkspaceWorldviewContent,
)

router = APIRouter()


# ════════════════════════════════════════════════════════════
# 导出辅助函数
# ════════════════════════════════════════════════════════════

def _markdown_to_plain_text(markdown: str) -> str:
    lines = []
    for line in markdown.splitlines():
        stripped = line.strip()
        if stripped.startswith("## "):
            lines.append(stripped[3:].strip())
        elif stripped.startswith("# "):
            lines.append(stripped[2:].strip())
        elif stripped == "---":
            lines.append("")
        else:
            lines.append(line)
    return "\n".join(lines).strip()


def _escape_pdf_text(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _safe_download_name(name: str, fallback: str, max_length: int = 80) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "", name).strip().strip(".")
    return (cleaned or fallback)[:max_length]


def _build_novel_docx(markdown: str, title: str) -> bytes:
    document = Document()
    document.add_heading(title, level=1)
    for block in _markdown_to_plain_text(markdown).split("\n\n"):
        text = block.strip()
        if not text or text == title:
            continue
        document.add_paragraph(text)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_novel_docx_zip(content: WorkspaceNovelChaptersContent) -> bytes:
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for index, chapter in enumerate(content.chapters, start=1):
            title = chapter.title.strip() or Path(chapter.filename).stem or f"chapter-{index:03d}"
            safe_title = _safe_download_name(title, f"chapter-{index:03d}")
            archive.writestr(f"{index:03d}-{safe_title}.docx", _build_novel_docx(chapter.markdown, title))
    return buffer.getvalue()


def _build_novel_pdf(markdown: str, title: str) -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4, leftMargin=54, rightMargin=54, topMargin=54, bottomMargin=54, title=title,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "NovelTitle", parent=styles["Title"], fontName="STSong-Light",
        fontSize=20, leading=28, spaceAfter=24,
    )
    chapter_style = ParagraphStyle(
        "ChapterTitle", parent=styles["Heading2"], fontName="STSong-Light",
        fontSize=15, leading=22, spaceBefore=8, spaceAfter=12,
    )
    body_style = ParagraphStyle(
        "NovelBody", parent=styles["BodyText"], fontName="STSong-Light",
        fontSize=11, leading=19, firstLineIndent=22, spaceAfter=7,
    )

    story = [Paragraph(_escape_pdf_text(title), title_style)]
    first_chapter = True
    for block in _markdown_to_plain_text(markdown).split("\n\n"):
        text = block.strip()
        if not text:
            continue
        if text.startswith("第") and "章" in text[:8]:
            if not first_chapter:
                story.append(PageBreak())
            first_chapter = False
            story.append(Paragraph(_escape_pdf_text(text), chapter_style))
            continue
        story.append(Paragraph(_escape_pdf_text(text).replace("\n", "<br/>"), body_style))
        story.append(Spacer(1, 4))

    doc.build(story)
    return buffer.getvalue()


# ════════════════════════════════════════════════════════════
# SSE watch 辅助
# ════════════════════════════════════════════════════════════

def _sse_event(event_type: str, payload: object) -> str:
    data = json.dumps(payload, ensure_ascii=False, default=str)
    return f"event: {event_type}\ndata: {data}\n\n"


def _classify_changes(changes, workspace_path: Path) -> set[str]:
    categories: set[str] = set()
    for _change_type, path_str in changes:
        try:
            rel = Path(path_str).relative_to(workspace_path)
        except ValueError:
            continue
        parts = rel.parts
        if not parts:
            continue
        top = parts[0]
        if top in ("outline.md", "evaluation.md"):
            categories.add("outline")
        elif top == "storyline.md" or (len(parts) > 1 and parts[0] == "storyline"):
            categories.add("storyline")
        elif top == "worldview.md":
            categories.add("worldview")
        elif len(parts) > 1 and parts[0] == "detail":
            categories.add("detail_outline")
        elif len(parts) > 1 and parts[0] == "character":
            categories.add("characters")
        elif len(parts) > 1 and parts[0] == "chapter":
            categories.add("novel")
        elif top == "novel.md":
            categories.add("novel")
    return categories


async def _workspace_watch_generator(owner_id: str, workspace_id: str, workspace_path: Path):
    thread_store = get_thread_store()
    _log("sse_open", channel="watch", workspace_id=workspace_id)
    start = time.perf_counter()
    try:
        async for changes in awatch(
            workspace_path,
            watch_filter=lambda _change, path: Path(path).suffix == ".md",
            debounce=400, step=50, recursive=True, ignore_permission_denied=True,
        ):
            categories = _classify_changes(changes, workspace_path)
            if not categories:
                continue
            if "outline" in categories:
                content = thread_store.artifacts.read_workspace_outline(owner_id, workspace_id)
                if content is not None:
                    yield _sse_event("outline", content.model_dump())
            if "storyline" in categories:
                content = thread_store.artifacts.read_workspace_storyline(owner_id, workspace_id)
                if content is not None:
                    yield _sse_event("storyline", content.model_dump())
            if "worldview" in categories:
                content = thread_store.artifacts.read_workspace_worldview(owner_id, workspace_id)
                if content is not None:
                    yield _sse_event("worldview", content.model_dump())
            if "detail_outline" in categories:
                content = thread_store.artifacts.read_workspace_detail_outline(owner_id, workspace_id)
                if content is not None:
                    yield _sse_event("detail_outline", content.model_dump())
            if "characters" in categories:
                content = thread_store.artifacts.read_workspace_characters(owner_id, workspace_id)
                if content is not None:
                    yield _sse_event("characters", content.model_dump())
            if "novel" in categories:
                content = thread_store.artifacts.read_workspace_novel_chapters(owner_id, workspace_id)
                if content is not None:
                    yield _sse_event("novel", content.model_dump())
        _log("sse_close", channel="watch", workspace_id=workspace_id,
             ms=int((time.perf_counter() - start) * 1000))
    except BaseException as exc:
        _log("sse_error", channel="watch", workspace_id=workspace_id,
             error=type(exc).__name__, ms=int((time.perf_counter() - start) * 1000))
        raise


# ════════════════════════════════════════════════════════════
# 端点
# ════════════════════════════════════════════════════════════

@router.get("/workspaces/{workspace_id}/bootstrap", response_model=WorkspaceBootstrapResponse)
def bootstrap_workspace(workspace_id: str, user: CurrentUser = Depends(current_user)) -> WorkspaceBootstrapResponse:
    thread_store = get_thread_store()
    data = thread_store.bootstrap_workspace(user.user_id, workspace_id)
    if data is None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    return WorkspaceBootstrapResponse(**data)


@router.get("/workspaces", response_model=list[WorkspaceSummary])
def list_workspaces(user: CurrentUser = Depends(current_user)) -> list[WorkspaceSummary]:
    return get_thread_store().list_workspaces(user.user_id)


@router.post("/workspaces", response_model=WorkspaceSummary)
def create_workspace(payload: WorkspaceCreateRequest, user: CurrentUser = Depends(current_user)) -> WorkspaceSummary:
    thread_store = get_thread_store()
    try:
        return thread_store.create_workspace(user.user_id, payload.title, payload.domain)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc


@router.get("/workspaces/{workspace_id}/outline", response_model=WorkspaceOutlineContent)
def get_workspace_outline(workspace_id: str, user: CurrentUser = Depends(current_user)) -> WorkspaceOutlineContent:
    content = get_thread_store().artifacts.read_workspace_outline(user.user_id, workspace_id)
    if content is None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    return content


@router.get("/workspaces/{workspace_id}/storyline", response_model=WorkspaceStorylineContent)
def get_workspace_storyline(workspace_id: str, user: CurrentUser = Depends(current_user)) -> WorkspaceStorylineContent:
    content = get_thread_store().artifacts.read_workspace_storyline(user.user_id, workspace_id)
    if content is None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    return content


@router.get("/workspaces/{workspace_id}/detail-outline", response_model=WorkspaceDetailOutlineContent)
def get_workspace_detail_outline(workspace_id: str, user: CurrentUser = Depends(current_user)) -> WorkspaceDetailOutlineContent:
    content = get_thread_store().artifacts.read_workspace_detail_outline(user.user_id, workspace_id)
    if content is None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    return content


@router.get("/workspaces/{workspace_id}/worldview", response_model=WorkspaceWorldviewContent)
def get_workspace_worldview(workspace_id: str, user: CurrentUser = Depends(current_user)) -> WorkspaceWorldviewContent:
    content = get_thread_store().artifacts.read_workspace_worldview(user.user_id, workspace_id)
    if content is None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    return content


@router.get("/workspaces/{workspace_id}/storyline-graph", response_model=WorkspaceStorylineGraphContent)
def get_workspace_storyline_graph(workspace_id: str, user: CurrentUser = Depends(current_user)) -> WorkspaceStorylineGraphContent:
    """故事线流程图。读取时按需生成兜底——图缺失/过期自动重生成（PR-03 上移到 API 层）。"""
    thread_store = get_thread_store()
    content = thread_store.artifacts.read_workspace_storyline_graph(user.user_id, workspace_id)
    if content is None:
        raise HTTPException(status_code=404, detail="Workspace not found")

    from app.domains.writing.expert_agent.services.storyline_graph import (
        build_storyline_graph_data, generate_storyline_graph, is_stale,
    )
    workspace = thread_store.get_workspace(user.user_id, workspace_id)
    if workspace is not None:
        ws_path = Path(workspace.workspace_path)
        if is_stale(ws_path):
            generate_storyline_graph(ws_path)
            content = thread_store.artifacts.read_workspace_storyline_graph(user.user_id, workspace_id)
            if content is None:
                raise HTTPException(status_code=404, detail="Workspace not found")
            content.stale = True
        data = build_storyline_graph_data(ws_path)
        if data is not None:
            content.storylines = [
                StorylineGraphStoryline(
                    id=sl.id, name=sl.name, type=sl.type, status=sl.status,
                    direction=sl.direction, key_events=sl.key_events,
                ) for sl in data.storylines
            ]
            content.events = {
                eid: StorylineGraphEvent(
                    id=ev.id, name=ev.name, type=ev.type,
                    storylines=list(ev.storylines), group=ev.group, doc_order=ev.doc_order,
                ) for eid, ev in data.events.items()
            }
            content.t_map = dict(data.t_map)
            content.storyline_count = len(data.storylines)
            content.event_count = len(data.events)
    return content


@router.get("/workspaces/{workspace_id}/characters", response_model=WorkspaceCharacterContent)
def get_workspace_characters(workspace_id: str, user: CurrentUser = Depends(current_user)) -> WorkspaceCharacterContent:
    content = get_thread_store().artifacts.read_workspace_characters(user.user_id, workspace_id)
    if content is None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    return content


@router.get("/workspaces/{workspace_id}/novel", response_model=WorkspaceNovelChaptersContent)
def get_workspace_novel(workspace_id: str, user: CurrentUser = Depends(current_user)) -> WorkspaceNovelChaptersContent:
    content = get_thread_store().artifacts.read_workspace_novel_chapters(user.user_id, workspace_id)
    if content is None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    return content


@router.get("/workspaces/{workspace_id}/watch")
async def watch_workspace(workspace_id: str, user: CurrentUser = Depends(current_user)):
    thread_store = get_thread_store()
    workspace = thread_store.get_workspace(user.user_id, workspace_id)
    if workspace is None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    workspace_path = Path(workspace.workspace_path)
    if not workspace_path.exists():
        raise HTTPException(status_code=404, detail="Workspace directory missing")
    return StreamingResponse(
        _workspace_watch_generator(user.user_id, workspace_id, workspace_path),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
    )


@router.get("/workspaces/{workspace_id}/novel/export.pdf")
def export_workspace_novel_pdf(workspace_id: str, user: CurrentUser = Depends(current_user)) -> Response:
    thread_store = get_thread_store()
    workspace = thread_store.get_workspace(user.user_id, workspace_id)
    if workspace is None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    content = thread_store.artifacts.read_workspace_novel(user.user_id, workspace_id)
    if content is None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    if not content.markdown.strip():
        raise HTTPException(status_code=404, detail="Novel content not found")
    filename = f"{workspace.title or workspace_id}.pdf"
    pdf = _build_novel_pdf(content.markdown, workspace.title or "小说正文")
    return Response(
        content=pdf, media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"},
    )


@router.get("/workspaces/{workspace_id}/novel/export-word.zip")
def export_workspace_novel_word_zip(workspace_id: str, user: CurrentUser = Depends(current_user)) -> Response:
    thread_store = get_thread_store()
    workspace = thread_store.get_workspace(user.user_id, workspace_id)
    if workspace is None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    content = thread_store.artifacts.read_workspace_novel_chapters(user.user_id, workspace_id)
    if content is None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    if not content.chapters:
        raise HTTPException(status_code=404, detail="Novel content not found")
    filename_base = _safe_download_name(workspace.title or workspace_id, workspace_id)
    archive = _build_novel_docx_zip(content)
    return Response(
        content=archive, media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(f'{filename_base}-word.zip')}"},
    )


@router.delete("/workspaces/{workspace_id}")
async def delete_workspace(workspace_id: str, user: CurrentUser = Depends(current_user)) -> dict[str, str | bool | list[str]]:
    thread_store = get_thread_store()
    try:
        deleted_thread_ids = thread_store.delete_workspace(user.user_id, workspace_id)
    except FileNotFoundError as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    if deleted_thread_ids is None:
        raise HTTPException(status_code=404, detail="Workspace not found")
    # T2.9：清理 checkpoint（分库）+ trace
    agent_service = get_agent_service()
    character_service = get_character_service()
    for thread_id in deleted_thread_ids:
        await agent_service.delete_thread_checkpoint(thread_id, owner_id=user.user_id)
        await character_service.delete_thread_checkpoint(thread_id)
    return {"status": "ok", "deleted": workspace_id, "deleted_threads": deleted_thread_ids}
